from docx import Document
from docx.shared import Inches
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
import pandas as pd
import os

def find_section_index(doc: Document, section_name: str):
    norm_target = section_name.strip().lower()
    for i, para in enumerate(doc.paragraphs):
        if para.style.name.startswith("Heading"):
            t = para.text.strip()
            if norm_target == t.lower():
                return i
    for i, para in enumerate(doc.paragraphs):
        if para.style.name.startswith("Heading"):
            if norm_target in para.text.strip().lower():
                return i
    return None

def insert_table(doc, idx, df):
    df = df.dropna(how='all')
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr_cells[i].text = str(col)
        hdr_cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for _, row in df.iterrows():
        if row.isnull().all(): continue
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = "" if pd.isna(val) else str(val)
            row_cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    tbl_idx = idx + 1
    doc._body._element.insert(tbl_idx, table._element)
    return table

def insert_picture(doc, idx, img_path, width_inches=5.5):
    new_para = doc.add_paragraph()
    doc._body._element.insert(idx+1, new_para._element)
    run = new_para.add_run()
    run.add_picture(img_path, width=Inches(width_inches))

def insert_text(doc, idx, content):
    new_para = doc.add_paragraph(content)
    doc._body._element.insert(idx+1, new_para._element)

def apply_actions(template_path: str, output_path: str, actions: list):
    doc = Document(template_path)
    errors = []
    for act in actions:
        section = act["target"]["section"]
        idx = find_section_index(doc, section)
        if idx is None:
            errors.append({"section": section, "error": f"Section '{section}' not found in template"})
            continue
        kind = act.get("kind")
        if kind == "text":
            for chunk in act["source"].get("content", []):
                insert_text(doc, idx, str(chunk))
        elif kind == "image":
            path = act["source"]["path"]
            try:
                if not os.path.exists(path):
                    errors.append({"section": section, "error": f"Image not found: {path}"})
                    continue
                insert_picture(doc, idx, path)
            except Exception as e:
                errors.append({"section": section, "error": f"Image failed: {e}"})
        elif kind == "table":
            try:
                wb = act["source"].get("workbook")
                sh = act["source"].get("sheet")
                df = None
                if wb and wb.endswith(('.xlsx', '.xls', '.xlsm')):
                    df = pd.read_excel(wb, sheet_name=sh)
                elif wb and wb.endswith('.csv'):
                    df = pd.read_csv(wb)
                if df is None or df.empty:
                    errors.append({"section": section, "error": f"No data in {wb}"})
                    continue
                insert_table(doc, idx, df)
            except Exception as e:
                errors.append({"section": section, "error": f"Table failed: {e}"})
        elif kind == "slides":
            slides_dir = act["source"].get("slides_dir")
            if slides_dir and os.path.isdir(slides_dir):
                for img in sorted(os.listdir(slides_dir)):
                    if img.lower().endswith(('.png', '.jpg', '.jpeg')):
                        insert_picture(doc, idx, os.path.join(slides_dir, img))
            else:
                errors.append({"section": section, "error": "Slides dir not found or no images."})
    doc.save(output_path)
    return {"errors": errors}


