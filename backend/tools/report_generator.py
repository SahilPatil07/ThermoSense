from pathlib import Path
from datetime import datetime
from typing import List, Dict, Any, Tuple, Optional
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import logging
from thefuzz import fuzz

# Import RAG system
try:
    from backend.tools.enhanced_rag_tools import rag_system
except ImportError:
    rag_system = None

logger = logging.getLogger(__name__)

def clean_xml_string(s: str) -> str:
    """Remove non-XML compatible characters (control characters, null bytes)"""
    if not isinstance(s, str):
        return str(s)
    # Remove null bytes and other control characters except \n, \r, \t
    return "".join(ch for ch in s if ch == '\n' or ch == '\r' or ch == '\t' or (ord(ch) >= 32 and ord(ch) != 127))

class ReportContextBuilder:
    """Builds comprehensive context for report generation"""
    
    def __init__(self, storage, session_id: str):
        self.storage = storage
        self.session_id = session_id
        
    def build_context(self) -> Dict[str, Any]:
        """Aggregate all available context"""
        memory = self.storage.load_memory(self.session_id)
        uploads = memory.get("uploads", [])
        chart_history = memory.get("chart_history", [])
        
        # Extract goal from chat history (simple heuristic for now)
        # In a real system, we might use LLM to summarize chat
        chat_history = self.storage.get_messages(self.session_id)
        project_goal = "Analyze thermal performance and identify anomalies."
        if chat_history:
            # Try to find user's first few messages
            for msg in chat_history[:3]:
                if msg["role"] == "user":
                    project_goal += f" User intent: {msg['content'][:100]}..."
                    
        # File metadata
        files_info = []
        for fname in uploads:
            fpath = self.storage.find_file(fname)
            if fpath:
                files_info.append({
                    "name": fname,
                    "size": fpath.stat().st_size
                })
                
        return {
            "session_id": self.session_id,
            "date": datetime.now().strftime("%Y-%m-%d"),
            "project_goal": project_goal,
            "files": files_info,
            "charts": chart_history,
            "chart_count": len(chart_history),
            "approved_charts": [c for c in chart_history if c.get("approved")],
            "data_analysis": memory.get("data_analysis", {})  # NEW: Include data analysis
        }

class SectionWriter:
    """Generates professional section content using LLM + RAG"""
    
    def __init__(self, llm_client):
        self.llm = llm_client
        
    def write_section(self, section_name: str, context: Dict[str, Any]) -> str:
        """Generate text for a specific section"""
        if not self.llm:
            return self._fallback_content(section_name)
            
        # Get section-specific detailed prompt
        prompt = self._get_section_prompt(section_name, context)
        
        try:
            response = self.llm.chat.completions.create(
                model="llama3.2",
                messages=[{"role": "user", "content": prompt}],
                temperature=0.1,
                max_tokens=300 # Slightly reduced
            )
            return response.choices[0].message.content.strip()
        except Exception as e:
            logger.error(f"Section generation failed: {e}")
            return self._fallback_content(section_name)
    
    def _get_section_prompt(self, section_name: str, context: Dict[str, Any]) -> str:
        """Get detailed section-specific prompts based on ER reference analysis"""
        
        # Extract context details
        files = [f['name'] for f in context.get('files', [])]
        file_info = ', '.join(files) if files else "thermal test data"
        chart_count = len(context.get('approved_charts', []))
        date = context.get('date', datetime.now().strftime('%Y-%m-%d'))
        
        # Retrieve style examples from RAG
        rag_example = ""
        if rag_system:
            rag_context, _, confidence, _ = rag_system.retrieve(
                f"Professional {section_name} section example from Engineering Report", 
                top_k=1
            )
            if rag_context and confidence > 0.3:
                rag_example = f"\n**STYLE REFERENCE (Match this tone and structure):**\n{rag_context[:500]}...\n"
        
        # Section-specific prompts based on reference ER analysis
        prompts = {
            "Executive Summary": f"""You are a Senior Automotive Thermal Engineer writing the Executive Summary for an official Engineering Report.

**DATA ANALYSIS RESULTS:**
{self._format_data_analysis(context.get('data_analysis', {}))}

**CONTEXT:**
- Test Date: {date}
- Files Analyzed: {file_info}
- Charts Generated: {chart_count}
{rag_example}

**WRITING REQUIREMENTS:**
Write a professional Executive Summary (4-5 paragraphs) that uses SPECIFIC VALUES from the data analysis:

1. **Test Overview** (1 paragraph):
   - State what test was conducted
   - Mention test configuration and conditions
   - Reference the data files analyzed

2. **Key Findings** (2-3 paragraphs):
   - Report SPECIFIC measured values from the data
   - Use actual temperature readings, flow rates, etc.
   - Mention which requirements were met or exceeded
   - Include margins where relevant
   - Example: "EM inlet temperature was above 50 degrees C during majority of test duration, with maximum recorded at [X] degrees C."

3. **Conclusion** (1 paragraph):
   - State overall assessment
   - Mention any critical issues or all-clear
   - Brief statement on thermal performance adequacy

**CRITICAL STYLE REQUIREMENTS:**
- NO asterisks or markdown formatting whatsoever
- NO bullet points - paragraphs only
- Use REAL VALUES from data analysis above
- Sound like a professional engineer, NOT AI-generated
- Use passive voice appropriately
- Be concise but include concrete numbers
- Temperature in degrees C, flow in lpm, pressure in bar

Write the Executive Summary now:""",

            "Introduction": f"""You are a Senior Automotive Thermal Engineer writing the Introduction section for an official Engineering Report.

**CONTEXT:**
- Test Date: {date}
- Data Files: {file_info}
- Analysis Scope: {chart_count} thermal performance charts generated
- Purpose: Evaluate thermal performance and identify potential issues
{rag_example}

**WRITING REQUIREMENTS:**
Write a professional Introduction (3-4 paragraphs) covering:

1. **Background & Context** (1 paragraph):
   - Briefly state the thermal system/component under test
   - Mention the test conditions or operational scenario
   - Example: "This report presents the thermal analysis performed on [system] under [conditions]. The evaluation was conducted to verify thermal performance and identify potential thermal management issues."

2. **Objectives** (1 paragraph):
   - State the main objective: thermal performance evaluation
   - Mention specific aspects: temperature distribution, heat transfer efficiency, cooling effectiveness
   - Example: "The primary objective is to assess the thermal characteristics, identify temperature hotspots, and evaluate the effectiveness of the cooling system."

3. **Scope** (1 paragraph):
   - Reference the data files analyzed
   - Mention the analysis methods (time-series analysis, correlation analysis)
   - Example: "The analysis is based on thermal test data recorded in {file_info}, encompassing temperature sensors, pressure measurements, and flow rates collected during [test duration/conditions]."

**STYLE GUIDELINES:**
- Use formal, technical language with passive voice where appropriate
- Be specific about thermal engineering aspects
- Avoid vague statements - use concrete terms (e.g., "temperature distribution" not "heat")
- NO markdown formatting (no #, **, etc.)
- Write in complete paragraphs, not bullet points

Write the Introduction now:""",

            "Objectives": f"""You are a Senior Automotive Thermal Engineer writing the Objectives section for an official Engineering Report.

**CONTEXT:**
- Test conducted: {date}
- Data analyzed: {file_info}
- Charts generated: {chart_count}
{rag_example}

**WRITING REQUIREMENTS:**
Write a professional Objectives section (2-3 paragraphs) that:

1. **Primary Objective** (1 paragraph):
   - Clearly state the main goal: thermal performance evaluation
   - Be specific about what thermal parameters are being assessed
   - Example: "The primary objective of this analysis is to evaluate the thermal performance of the system by analyzing temperature profiles, heat transfer rates, and thermal stability under operational conditions."

2. **Specific Sub-objectives** (1-2 paragraphs):
   - Identify temperature hotspots and thermal gradients
   - Assess cooling system effectiveness
   - Verify thermal design margins
   - Detect thermal anomalies or abnormal behavior
   - Example: "Specifically, the analysis aims to: (a) identify critical temperature zones and assess thermal gradients across components, (b) evaluate the cooling system's heat dissipation capability, and (c) detect any thermal anomalies that may indicate design issues."

3. **Expected Outcomes**:
   - Mention what deliverables are expected (charts, analysis, recommendations)
   - Example: "The expected outcome is a comprehensive thermal assessment with detailed visualizations and actionable recommendations for thermal management optimization."

**STYLE GUIDELINES:**
- Use precise thermal engineering terminology
- Structure with clear, numbered objectives if multiple sub-objectives
- Be measurable and specific
- NO markdown or bullet points - write in paragraph form

Write the Objectives section now:""",

            "Requirements": f"""You are a Senior Automotive Thermal Engineer writing the Requirements section for an official Engineering Report.

**CONTEXT:**
- System analyzed: Based on {file_info}
- Test parameters measured: Temperature, Pressure, Flow (typical thermal parameters)
{rag_example}

**WRITING REQUIREMENTS:**
Write a professional Requirements section (2-3 paragraphs) covering:

1. **Operational Requirements** (1 paragraph):
   - Define normal operating temperature ranges
   - Specify acceptable thermal performance criteria
   - Example: "The thermal system shall maintain component temperatures within the specified operational range of [X]°C to [Y]°C under normal operating conditions. Cooling system efficiency shall be sufficient to prevent thermal runaway or excessive temperature gradients."

2. **Design Requirements** (1 paragraph):
   - Thermal safety margins
   - Heat transfer requirements
   - Temperature uniformity requirements
   - Example: "Design requirements stipulate a minimum [X]°C safety margin below critical temperature thresholds, with thermal gradients not exceeding [Y]°C across critical components to ensure structural integrity and prevent thermal stress."

3. **Test/Data Requirements** (1 paragraph):
   - Measurement accuracy and sensor requirements
   - Data collection methods
   - Example: "Thermal data shall be collected using calibrated temperature sensors with ±0.5°C accuracy, sampled at minimum [frequency] to capture thermal transients and steady-state behavior."

**STYLE GUIDELINES:**
- Use "shall" for mandatory requirements (standard engineering practice)
- Be specific with numerical ranges where possible (even if estimated)
- Write in formal, specification-style language
- NO markdown formatting

Write the Requirements section now:""",

            "Conclusion": f"""You are a Senior Automotive Thermal Engineer writing the Conclusion section for an official Engineering Report.

**CONTEXT:**
- Data analyzed: {file_info}
- Number of charts/analyses: {chart_count}
- Test date: {date}
{rag_example}

**WRITING REQUIREMENTS:**
Write a professional Conclusion section (2-3 paragraphs) that:

1. **Summary of Findings** (1 paragraph):
   - Recap the main thermal performance observations
   - Mention key temperature trends identified
   - Reference critical findings from the analysis
   - Example: "The thermal analysis of {file_info} reveals that the system operates within acceptable temperature ranges under normal conditions. Temperature profiles show [expected/stable/variable] behavior with [trend description]."

2. **Assessment** (1 paragraph):
   - State whether thermal performance meets requirements
   - Highlight any concerns or anomalies discovered
   - Assess cooling system effectiveness
   - Example: "Overall thermal performance is [satisfactory/requires attention/critical], with [X number] of thermal parameters within specification. However, [mention any concerns like hotspots, excessive gradients, cooling inefficiencies]."

3. **Forward-Looking Statement** (1 paragraph):
   - Briefly mention next steps or recommendations
   - Indicate if further analysis or testing is needed
   - Example: "Based on these findings, [continued monitoring/design optimization/further testing] is recommended to [maintain thermal stability/improve cooling efficiency/address identified issues]."

**STYLE GUIDELINES:**
- Be definitive but measured in conclusions
- Use evidence-based statements referencing the data analyzed
- Maintain professional, objective tone
- NO markdown formatting
- Write in complete paragraphs

Write the Conclusion section now:"""
        }
        
        # Return section-specific prompt or generic if not defined
        return prompts.get(section_name, self._generic_prompt(section_name, context, rag_example))
    
    def _format_data_analysis(self, data_analysis: Dict) -> str:
        """Format data analysis results for prompt"""
        if not data_analysis or 'key_findings' not in data_analysis:
            return "No detailed data analysis available."
        
        findings = data_analysis.get('key_findings', [])
        summary = data_analysis.get('executive_summary', '')
        
        formatted = "Key Findings from Data:\n"
        for i, finding in enumerate(findings[:8], 1):
            formatted += f"{i}. {finding}\n"
        
        if summary:
            formatted += f"\nPre-generated summary:\n{summary}"
        
        return formatted
    
    def _generic_prompt(self, section_name: str, context: Dict, rag_example: str) -> str:
        """Generic prompt for undefined sections"""
        files = [f['name'] for f in context.get('files', [])]
        return f"""You are a Senior Engineer writing the {section_name} section for a technical Engineering Report.

**CONTEXT:**
- Files analyzed: {', '.join(files) if files else 'thermal test data'}
- Date: {context.get('date')}
{rag_example}

Write a professional {section_name} section (2-3 paragraphs) with:
- Technical depth appropriate for engineering documentation
- Formal language and passive voice where suitable
- Specific references to thermal engineering concepts
- NO markdown formatting

Write the {section_name} section now:"""

    def _fallback_content(self, section_name: str) -> str:
        defaults = {
            "Introduction": "This report presents a comprehensive thermal analysis performed on the provided datasets. The evaluation aims to assess thermal performance, identify temperature anomalies, and verify that the system operates within specified thermal limits under test conditions.",
            "Objectives":  "The primary objective is to evaluate the thermal characteristics and performance of the system. Specific goals include identifying critical temperature zones, assessing cooling system effectiveness, and detecting thermal anomalies that may indicate design or operational issues.",
            "Requirements": "The thermal system shall maintain component temperatures within specified operational ranges under normal conditions. Design requirements stipulate adequate thermal margins and heat dissipation capability to prevent thermal runaway. Data collection shall employ calibrated sensors with sufficient accuracy and sampling rates.",
            "Conclusion": "The thermal analysis indicates that the system demonstrates acceptable thermal performance under the test conditions evaluated. Temperature profiles exhibit expected behavior patterns. Continued monitoring and further analysis are recommended to ensure sustained thermal integrity."
        }
        return defaults.get(section_name, f"Professional technical content for {section_name} section.")

class ThermalReportGenerator:
    """Generate professional thermal analysis reports using corporate template"""
    
    def __init__(self, llm_client=None):
        self.llm = llm_client
        self.template_path = Path("assets/template.docx")
        self.section_writer = SectionWriter(llm_client)
    
    def generate_report(
        self,
        session_id: str,
        storage,
        output_path: Path
    ) -> Tuple[bool, str, str]:
        """
        Generate report by populating the corporate template
        """
        try:
            if not self.template_path.exists():
                return False, f"Template not found at {self.template_path}", ""
            
            # Load template
            doc = Document(self.template_path)
            
            # Build Context
            context_builder = ReportContextBuilder(storage, session_id)
            context = context_builder.build_context()
            
            # Map our sections to template headers
            # Template headers found: Summary, Introduction, Results, Discussion, Conclusion
            section_map = {
                "Executive Summary": "Summary",
                "Introduction": "Introduction",
                "Methodology": "Introduction",
                "Objectives": "Objectives",
                "Test Objects": "Test Objects",
                "Requirements": "Requirements",
                "Results": "Results",
                "Analysis and Results": "Results",
                "Discussion": "Discussion",
                "Conclusion": "Conclusion",
                "Recommendations": "Conclusion",
                "Recommendation": "Conclusion",
                "Concluding Remarks": "Conclusion",
                "Summary": "Summary",
                "Background": "Introduction",
                "Findings": "Results",
                "Analysis": "Results",
                "Interpretation": "Discussion",
                "Appendices": "Conclusion"
            }
            
            # Track insertion points to avoid messing up the document order
            # We'll do a pass to find all headers first
            headers = self._find_headers(doc)
            
            # Get existing report content
            report_content = storage.get_report_content(session_id)
            session_dir = storage.get_session_dir(session_id)
            
            # AUTO-POPULATION: Add approved charts that aren't already in the report
            # Build set of existing chart IDs across all sections
            existing_chart_ids = set()
            for section_items in report_content.values():
                if section_items:
                    for item in section_items:
                        if item.get("type") == "chart" and isinstance(item.get("content"), dict):
                            chart_id = item["content"].get("chart_id")
                            if chart_id:
                                existing_chart_ids.add(chart_id)
            
            # Auto-populate missing approved charts with FULL metadata
            approved_charts = context["approved_charts"]
            if approved_charts:
                if "Analysis and Results" not in report_content:
                    report_content["Analysis and Results"] = []
                    
                charts_added = 0
                for chart in approved_charts:
                    chart_id = chart.get("chart_id")
                    # Only add if not already in report
                    if chart_id and chart_id not in existing_chart_ids:
                        # CRITICAL FIX: Include ALL chart metadata for proper insertion
                        report_content["Analysis and Results"].append({
                            "type": "chart",
                            "content": {
                                "chart_id": chart_id,
                                "path": chart.get("chart_path") or chart.get("path"),
                                "x_column": chart.get("x_column", ""),
                                "y_columns": chart.get("y_columns", []),
                                "chart_type": chart.get("chart_type", "line"),
                                "summary": chart.get("summary", "")
                            },
                            "id": f"chart_{chart_id}"  # Add unique ID for deduplication
                        })
                        existing_chart_ids.add(chart_id)
                        charts_added += 1
                
                if charts_added > 0:
                    print(f"Auto-populated {charts_added} new charts into report")
            
            # DYNAMIC CONTENT GENERATION - DISABLED as per user request
            # Speed optimization: Only generate Executive Summary and Conclusion if missing
            # Other sections can be empty or use fallback
            # critical_sections = ["Executive Summary", "Conclusion"]
            # for section in critical_sections:
            #     if not report_content.get(section):
            #         print(f"Generating {section} using SectionWriter...")
            #         text = self.section_writer.write_section(section, context)
            #         report_content[section] = [{"type": "text", "content": text}]
            
            # Iterate through our content and insert into template
            total_items = 0
            for section_name, items in report_content.items():
                if not items:
                    continue
                    
                target_header = section_map.get(section_name, section_name)
                target_paragraph = headers.get(target_header)
                
                if target_paragraph:
                    # Insert content after this header
                    # Deduplicate items by ID if possible
                    unique_items = []
                    seen_ids = set()
                    for item in items:
                        item_id = item.get("id")
                        # Special check for charts to avoid duplicates
                        if item.get("type") == "chart" and isinstance(item.get("content"), dict):
                            chart_id = item["content"].get("chart_id")
                            if chart_id:
                                if chart_id in seen_ids:
                                    continue
                                seen_ids.add(chart_id)
                        
                        if item_id and item_id not in seen_ids:
                            seen_ids.add(item_id)
                            unique_items.append(item)
                        elif not item_id:
                            unique_items.append(item)
                            
                    self._insert_content_after(doc, target_paragraph, unique_items, session_dir, storage, session_id)
                    total_items += len(unique_items)
                else:
                    print(f"Header '{target_header}' not found in template. Appending to end.")
                    doc.add_heading(section_name, level=2)
                    self._insert_content_after(doc, doc.paragraphs[-1], items, session_dir, storage, session_id)
                    total_items += len(items)
            
            # Save with retry mechanism for PermissionError
            max_retries = 3
            for attempt in range(max_retries):
                try:
                    doc.save(output_path)
                    logger.info(f"Report saved successfully to {output_path}")
                    break
                except PermissionError:
                    if attempt < max_retries - 1:
                        import time
                        # Try a different filename if it's locked
                        timestamp = int(time.time())
                        new_path = output_path.parent / f"thermal_report_{timestamp}.docx"
                        logger.warning(f"Permission denied for {output_path.name}. Retrying with {new_path.name}")
                        output_path = new_path
                    else:
                        raise
            
            return True, f"Report generated with {total_items} items", str(output_path)
        
        except Exception as e:
            import traceback
            traceback.print_exc()
            return False, f"Report generation failed: {str(e)}", ""

    def _find_headers(self, doc: Document) -> Dict[str, Any]:
        """Find paragraph objects for each known header using fuzzy matching"""
        headers = {}
        # Expanded list of target headers based on common engineering report structures
        targets = [
            "Summary", "Executive Summary", 
            "Introduction", "Background", "Objectives",
            "Test Objects", "Requirements",
            "Results", "Analysis and Results", "Findings",
            "Discussion", "Analysis", "Interpretation",
            "Conclusion", "Recommendations", "Recommendation", "Summary and Conclusion", "Concluding"
        ]
        
        # Map targets to normalized keys for easier lookup
        normalized_map = {
            "Summary": "Summary",
            "Executive Summary": "Summary",
            "Introduction": "Introduction",
            "Background": "Introduction",
            "Objectives": "Objectives",
            "Test Objects": "Test Objects",
            "Requirements": "Requirements",
            "Results": "Results",
            "Analysis and Results": "Results",
            "Findings": "Results",
            "Discussion": "Discussion",
            "Analysis": "Discussion",
            "Interpretation": "Discussion",
            "Conclusion": "Conclusion",
            "Recommendations": "Conclusion",
            "Recommendation": "Conclusion",
            "Summary and Conclusion": "Conclusion",
            "Concluding": "Conclusion"
        }
        
        for p in doc.paragraphs:
            text = p.text.strip()
            if not text or len(text) > 60: # Headers are usually short
                continue
                
            # Try fuzzy matching
            for t in targets:
                # Use ratio for fuzzy match (threshold 80)
                score = fuzz.partial_ratio(t.lower(), text.lower())
                if score > 80:
                    norm_key = normalized_map.get(t)
                    if norm_key not in headers: # Keep the first match
                        headers[norm_key] = p
                        logger.info(f"Matched header '{text}' to '{norm_key}' (target: {t}, score: {score})")
                    break
        
        # If some standard headers are missing, try to find ANY heading-like paragraph
        # This is a fallback for non-standard templates
        if not headers:
            logger.warning("No standard headers found. Attempting to find any bold/large paragraphs.")
            for p in doc.paragraphs:
                if p.style.name.startswith('Heading') or any(r.bold for r in p.runs):
                    text = p.text.strip()
                    if text and len(text) < 50:
                        headers[text] = p
                        
        return headers

    def _insert_content_after(self, doc: Document, anchor_p, items: List[Dict], session_dir: Path, storage, session_id: str):
        """Insert content items immediately after the anchor paragraph"""
        # We need to insert elements *after* the anchor. 
        # python-docx doesn't make "insert after" easy, so we insert *before* the *next* element.
        # But determining the next element is tricky.
        # A simpler approach for this task: 
        # 1. Find the index of the anchor paragraph
        # 2. Insert new paragraphs after it.
        
        # However, doc.paragraphs is a list of proxies. 
        # The most robust way in python-docx to "insert after" is to call `insert_paragraph_before` 
        # on the *next* paragraph.
        
        try:
            # Find index
            p_element = anchor_p._element
            parent = p_element.getparent()
            index = parent.index(p_element) + 1
            
            for item in reversed(items): # Reverse because we insert at the same position repeatedly
                item_type = item.get("type")
                content = item.get("content")
                
                # We will create new paragraphs/runs and move them to the correct position
                
                if item_type == "text":
                    new_p = doc.add_paragraph(clean_xml_string(content))
                    self._move_p_after(parent, new_p, index)
                    
                elif item_type == "chart":
                    # Handle chart
                    chart_data = content
                    if "x_column" not in content and "chart_id" in content:
                        chart_data = self._find_chart_in_history(content["chart_id"], storage, session_id)
                    
                    if chart_data:
                        # Try multiple paths to find the chart image
                        chart_id = chart_data.get('chart_id')
                        
                        # Option 1: Use path from chart_data if provided
                        if chart_data.get('path'):
                            chart_path = Path(chart_data['path'])
                        else:
                            # Option 2: Standard location
                            chart_path = session_dir / f"chart_{chart_id}.png"
                        
                        # Option 3: If not found, try chart_path from history
                        if not chart_path.exists() and chart_data.get('chart_path'):
                            chart_path = Path(chart_data['chart_path'])
                        
                        # Debug logging
                        print(f"Inserting chart {chart_id}")
                        print(f"   Session dir: {session_dir}")
                        print(f"   Chart path: {chart_path}")
                        print(f"   Absolute path: {chart_path.absolute()}")
                        print(f"   Exists: {chart_path.exists()}")
                        
                        if chart_path.exists():
                            # Image
                            img_p = doc.add_paragraph()
                            img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run = img_p.add_run()
                            run.add_picture(str(chart_path.absolute()), width=Inches(5.5))
                            self._move_p_after(parent, img_p, index)
                            
                            # Caption
                            y_cols = chart_data.get('y_columns', [])
                            if isinstance(y_cols, list):
                                y_label = ', '.join(y_cols)
                            else:
                                y_label = str(y_cols)
                                
                            cap_p = doc.add_paragraph(f"Figure: {y_label} vs {chart_data.get('x_column', 'X')}")
                            cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            self._move_p_after(parent, cap_p, index)
                            
                            # Analysis
                            analysis = self._generate_chart_analysis(chart_data, storage)
                            ana_p = doc.add_paragraph(analysis)
                            self._move_p_after(parent, ana_p, index)
                            print(f"   Chart inserted successfully")
                        else:
                            print(f"   Chart file not found: {chart_path}")

                elif item_type == "image":
                     content_path = Path(content)
                     if content_path.exists():
                        if content_path.suffix.lower() == '.txt':
                            # Handle text fallback for slides
                            try:
                                with open(content_path, 'r', encoding='utf-8') as f:
                                    text_content = f.read()
                                # Add title if it looks like a slide
                                if "Slide" in text_content and ":" in text_content.split('\n')[0]:
                                    title_line = text_content.split('\n')[0]
                                    p = doc.add_paragraph()
                                    run = p.add_run(clean_xml_string(title_line))
                                    run.bold = True
                                    self._move_p_after(parent, p, index)
                                    text_to_add = '\n'.join(text_content.split('\n')[1:])
                                else:
                                    text_to_add = text_content
                                
                                new_p = doc.add_paragraph(clean_xml_string(text_to_add))
                                self._move_p_after(parent, new_p, index)
                                logger.info(f"Inserted text fallback for slide from {content_path.name}")
                            except Exception as e:
                                logger.error(f"Failed to read text fallback slide: {e}")
                        else:
                            # Standard image
                            img_p = doc.add_paragraph()
                            img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run = img_p.add_run()
                            # Ensure path is absolute and string
                            img_path = str(content_path.absolute())
                            try:
                                run.add_picture(img_path, width=Inches(5.5))
                                self._move_p_after(parent, img_p, index)
                                logger.info(f"Inserted image from {content_path.name}")
                            except Exception as e:
                                logger.error(f"Failed to insert image {img_path}: {e}")

                elif item_type == "table":
                    # Handle table
                    logger.info(f"Processing table insertion. Content type: {type(content)}")
                    
                    if isinstance(content, str):
                        # Fallback: if content is a string (e.g. markdown), insert as text
                        logger.warning("Table content is a string, inserting as text paragraph.")
                        new_p = doc.add_paragraph(clean_xml_string(content))
                        self._move_p_after(parent, new_p, index)
                        continue

                    table_data = content.get("data", [])
                    columns = content.get("columns", [])
                    
                    if table_data and columns:
                        # Create table
                        table = doc.add_table(rows=1, cols=len(columns))
                        try:
                            table.style = 'Table Grid'
                        except:
                            logger.warning("Style 'Table Grid' not found in template, using default.")
                        
                        # Header
                        hdr_cells = table.rows[0].cells
                        for i, col in enumerate(columns):
                            hdr_cells[i].text = clean_xml_string(str(col))
                            # Make bold
                            for paragraph in hdr_cells[i].paragraphs:
                                for run in paragraph.runs:
                                    run.font.bold = True
                        
                        # Rows
                        for row_data in table_data:
                            row_cells = table.add_row().cells
                            for i, col in enumerate(columns):
                                val = row_data.get(col, "")
                                row_cells[i].text = clean_xml_string(str(val))
                        
                        # Apply borders manually to ensure they are visible
                        self._apply_table_borders(table)
                        
                        # Move table using the same helper (it works for any element)
                        self._move_p_after(parent, table, index)
                        logger.info(f"Table inserted successfully with {len(table_data)} rows")
                    else:
                        logger.warning(f"Table data or columns missing. Data: {len(table_data)}, Columns: {len(columns)}")
        except Exception as e:
            print(f"Error inserting content: {e}")

    
    def _move_p_after(self, parent, element, index):
        """Move a paragraph or table element to a specific index in the parent"""
        parent.insert(index, element._element)

    def _apply_table_borders(self, table):
        """Manually apply borders to a table using XML manipulation"""
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        
        tbl = table._element
        tblPr = tbl.xpath('w:tblPr')[0]
        
        # Create tblBorders element
        tblBorders = OxmlElement('w:tblBorders')
        
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4') # 1/2 pt
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            tblBorders.append(border)
            
        tblPr.append(tblBorders)

    def _find_chart_in_history(self, chart_id: str, storage, session_id: str) -> Optional[Dict]:
        """Look up chart details from history"""
        try:
            memory = storage.load_memory(session_id)
            for chart in memory.get("chart_history", []):
                if chart.get("chart_id") == chart_id:
                    return chart
        except:
            pass
        return None

    def _generate_chart_analysis(self, chart: Dict, storage) -> str:
        """Generate AI analysis for chart with specific anomaly data points"""
        if self.llm:
            try:
                # Handle Heatmaps vs Charts
                if chart.get('chart_type') == 'heatmap':
                    prompt = f"""Analyze this correlation heatmap for a report:
                    Variables: {', '.join(chart.get('y_columns', []))}
                    
                    Analysis Summary: {chart.get('summary', 'N/A')}
                    
                    Provide a 2-sentence technical summary of the correlations observed."""
                else:
                    # Get anomaly details if available
                    stats = chart.get('stats', {})
                    anomaly_info = ""
                    if isinstance(stats, dict) and 'anomalies' in stats:
                        anomaly_info = f"\nAnomalies detected: {stats['anomalies']}"
                    
                    prompt = f"""Analyze this {chart.get('chart_type', 'chart')} for a  thermal analysis report.
                    Chart: {', '.join(chart.get('y_columns', []))} vs {chart.get('x_column')}
                    Data Summary: {chart.get('summary', 'N/A')}
                    Statistics: {stats}{anomaly_info}
                    
                    Provide a concise 2-sentence summary mentioning key trends and any anomalies with specific data point values"""
                
                response = self.llm.chat.completions.create(
                    model="llama3.2",
                    messages=[{"role": "user", "content": prompt}],
                    temperature=0.1,
                    max_tokens=80  # Further reduced for speed
                )
                return response.choices[0].message.content.strip()
            except:
                pass
        return f"Chart showing {chart.get('y_columns')} vs {chart.get('x_column')}."

# Singleton
_report_generator = None

def get_report_generator(llm_client=None):
    global _report_generator
    if _report_generator is None:
        _report_generator = ThermalReportGenerator(llm_client)
    return _report_generator 

